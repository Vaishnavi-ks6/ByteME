#!/usr/bin/env python3
import os
from docx import Document
from docx.shared import Inches
from docx2pdf import convert

def build_document() -> Document:
    doc = Document()

    # Title
    doc.add_heading('3. Methodology', level=1)

    # Section A
    doc.add_heading('A. Data Collection and Preprocessing', level=2)
    doc.add_paragraph("""hsre""")

    # Section B
    doc.add_heading('B. Feature Selection', level=2)
    doc.add_paragraph("hsre")
    doc.add_paragraph("Selected Features:")
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'Unit'
    features = [
        ('Daily temperature', '°C'),
        ('Relative humidity', '%'),
        ('Soil moisture content', '%'),
        ('Rainfall', 'mm'),
        ('Wind speed', 'm/s'),
        ('Solar radiation', 'MJ/m²'),
        ('Crop growth stage', '—')
    ]
    for feature, unit in features:
        row_cells = table.add_row().cells
        row_cells[0].text = feature
        row_cells[1].text = unit

    # Section C
    doc.add_heading('C. Model Selection and Training', level=2)
    doc.add_paragraph("""hsre""")

    # Section D
    doc.add_heading('D. Model Evaluation and Performance Metrics', level=2)
    doc.add_paragraph("""hsre""")

    # Section E
    doc.add_heading('E. Visualization and Interpretation', level=2)
    doc.add_paragraph("""hsre""")

    # Section F
    doc.add_heading('F. Deployment and Future Enhancements', level=2)
    doc.add_paragraph("""hsre""")

    # Section 4
    doc.add_heading('4. Results and Discussion', level=1)

    # A. Correlation Analysis
    doc.add_heading('A. Correlation Analysis', level=2)
    doc.add_paragraph("""hsre""")

    # B. Model Performance Comparison
    doc.add_heading('B. Model Performance Comparison', level=2)
    doc.add_paragraph("hsre")
    perf_table = doc.add_table(rows=1, cols=4)
    hdr = perf_table.rows[0].cells
    hdr[0].text = 'Model'
    hdr[1].text = 'RMSE'
    hdr[2].text = 'MAE'
    hdr[3].text = 'R² Score'
    data = [
        ('Linear Regression', '47.32', '35.21', '0.82'),
        ('Decision Tree', '52.15', '38.67', '0.78')
    ]
    for model, rmse, mae, r2 in data:
        row = perf_table.add_row().cells
        row[0].text = model
        row[1].text = rmse
        row[2].text = mae
        row[3].text = r2

    doc.add_paragraph("""hsre""")

    return doc

def main():
    # Build and save DOCX
    doc = build_document()
    docx_path = "Formatted_Methodology.docx"

    # Save document and explicitly close it
    doc.save(docx_path)
    del doc  # Explicitly delete the document object to release file handles

    print(f"Saved Word file at: {os.path.abspath(docx_path)}")

    # Convert to PDF
    pdf_path = "Formatted_Methodology.pdf"

    # Add a small delay to ensure the file is fully written and closed
    import time
    time.sleep(1)

    # Convert to PDF
    convert(docx_path, pdf_path)

    # Add another small delay to ensure PDF is fully written
    time.sleep(1)

    print(f"Saved PDF file at:  {os.path.abspath(pdf_path)}")

if __name__ == "__main__":
    main()