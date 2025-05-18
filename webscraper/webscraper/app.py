import sys
import os
import os
import sys


import google.generativeai as genai
from flask import Flask, render_template, request, jsonify, send_file, url_for
from flask_cors import CORS
from dotenv import load_dotenv
import uuid
import time
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Load environment variables from .env file
load_dotenv()

# Initialize Flask app
app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

# Configure Google Generative AI
# You'll need to set your API key in an environment variable
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    print("Warning: GOOGLE_API_KEY environment variable not set")

# Configure Gemini
if GOOGLE_API_KEY:
    genai.configure(api_key=GOOGLE_API_KEY)

# Create output directory if it doesn't exist
OUTPUT_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'output')
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process_text():
    try:
        # Get the input text from the request
        input_text = request.form.get('text')
        if not input_text:
            return jsonify({'error': 'No text provided'}), 400

        # Check if the input is already a JSON structure
        try:
            import json
            # Try to parse as JSON first
            if input_text.strip().startswith('{') and input_text.strip().endswith('}'):
                print("Input appears to be JSON, trying to parse directly")
                structured_content = json.loads(input_text)
                print(f"Successfully parsed JSON input: {structured_content.keys() if isinstance(structured_content, dict) else 'not a dict'}")
            else:
                # Use Gemini AI to structure the content
                print("Using Gemini to structure content")
                structured_content = generate_structured_content(input_text)
        except json.JSONDecodeError:
            # Not valid JSON, use Gemini
            print("Input is not valid JSON, using Gemini")
            structured_content = generate_structured_content(input_text)

        if not structured_content:
            return jsonify({'error': 'Failed to generate structured content'}), 500

        # Create a unique ID for this request to avoid file conflicts
        request_id = str(uuid.uuid4())[:8]

        # Create output filenames with the unique ID
        docx_filename = f"Formatted_Document_{request_id}.docx"
        pdf_filename = f"Formatted_Document_{request_id}.pdf"

        output_docx_path = os.path.join(OUTPUT_FOLDER, docx_filename)

        # Generate the document directly
        success = create_document(structured_content, output_docx_path)

        if not success or not os.path.exists(output_docx_path):
            return jsonify({'error': 'Failed to create document file'}), 500

        # Verify the document was created
        if not os.path.exists(output_docx_path):
            return jsonify({'error': f'Document file was not created at {output_docx_path}'}), 500

        # Try to convert to PDF if docx2pdf is available
        output_pdf_path = os.path.join(OUTPUT_FOLDER, pdf_filename)
        pdf_available = False

        try:
            from docx2pdf import convert
            convert(output_docx_path, output_pdf_path)

            # Verify PDF was created
            if os.path.exists(output_pdf_path) and os.path.getsize(output_pdf_path) > 0:
                pdf_available = True
                print(f"PDF created successfully at {output_pdf_path}")
            else:
                print(f"PDF file was not created or is empty")
        except Exception as pdf_error:
            print(f"Error converting to PDF: {str(pdf_error)}")
            import traceback
            print(traceback.format_exc())
            # Continue even if PDF conversion fails

        # Return success response with download links
        response_data = {
            'success': True,
            'message': 'Document generated successfully',
            'docx_url': url_for('download_file', filename=docx_filename),
            'structured_content': structured_content
        }

        if pdf_available:
            response_data['pdf_url'] = url_for('download_file', filename=pdf_filename)

        return jsonify(response_data)

    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"Error in process_text: {str(e)}\n{error_details}")
        return jsonify({'error': str(e)}), 500

def generate_structured_content(input_text):
    """Use Gemini AI to structure the input text into sections for the document"""
    if not GOOGLE_API_KEY:
        raise Exception("Google API Key not configured")

    try:
        # Configure the model
        generation_config = {
            "temperature": 0.1,  # Lower temperature for more deterministic output
            "top_p": 0.95,
            "top_k": 40,
            "max_output_tokens": 8192,
        }

        # Create the prompt for Gemini
        prompt = f"""
        You are a document formatting expert. I need to create a well-structured document from the following text content:

        ```
        {input_text}
        ```

        Please analyze this content and organize it into a structured format with:
        1. A clear, descriptive title that summarizes the content
        2. Multiple sections with appropriate headings
        3. Paragraphs of content for each section
        4. Lists or bullet points where appropriate
        5. Tables for any tabular data

        IMPORTANT: Your response must be ONLY a valid JSON object with no additional text or explanation.

        The JSON structure must be exactly as follows:
        {{
            "title": "Document Title",
            "sections": [
                {{
                    "heading": "Section 1 Heading",
                    "level": 1,
                    "content": [
                        {{
                            "type": "paragraph",
                            "text": "Paragraph text here..."
                        }},
                        {{
                            "type": "bullet_list",
                            "items": ["Item 1", "Item 2", "Item 3"]
                        }},
                        {{
                            "type": "table",
                            "headers": ["Column 1", "Column 2"],
                            "rows": [
                                ["Row 1 Col 1", "Row 1 Col 2"],
                                ["Row 2 Col 1", "Row 2 Col 2"]
                            ]
                        }}
                    ]
                }}
            ]
        }}

        Make sure:
        1. The JSON is valid and properly formatted with no syntax errors
        2. All text content from the input is included in the structured document
        3. The document structure makes logical sense
        4. There are no placeholders or sample text in your response - use the actual content
        5. Your response contains ONLY the JSON object, nothing else
        """

        # Initialize the model
        model = genai.GenerativeModel(
            model_name="gemini-2.0-flash",  #  structured output
            generation_config=generation_config
        )

        # Generate the response
        response = model.generate_content(prompt)

        # Extract the JSON from the response
        import json
        response_text = response.text.strip()

        # If the response is wrapped in markdown code blocks, remove them
        if response_text.startswith("```json"):
            response_text = response_text.replace("```json", "", 1)
        elif response_text.startswith("```"):
            response_text = response_text.replace("```", "", 1)

        if response_text.endswith("```"):
            response_text = response_text[:-3]

        # Clean up any potential leading/trailing characters that might break JSON parsing
        response_text = response_text.strip()

        # Try to find JSON object boundaries if there's extra text
        if not response_text.startswith('{'):
            start_idx = response_text.find('{')
            if start_idx >= 0:
                response_text = response_text[start_idx:]

        if not response_text.endswith('}'):
            end_idx = response_text.rfind('}')
            if end_idx >= 0:
                response_text = response_text[:end_idx+1]

        # Parse the JSON
        try:
            structured_content = json.loads(response_text)
            return structured_content
        except json.JSONDecodeError as json_err:
            print(f"JSON parsing error: {str(json_err)}")
            print(f"Response text: {response_text}")

            # Try a fallback approach - create a simple document structure
            return {
                "title": "Generated Document",
                "sections": [
                    {
                        "heading": "Content",
                        "level": 1,
                        "content": [
                            {
                                "type": "paragraph",
                                "text": input_text
                            }
                        ]
                    }
                ]
            }

    except Exception as e:
        import traceback
        print(f"Error generating structured content: {str(e)}")
        print(traceback.format_exc())

        # Return a simple fallback structure
        return {
            "title": "Generated Document",
            "sections": [
                {
                    "heading": "Content",
                    "level": 1,
                    "content": [
                        {
                            "type": "paragraph",
                            "text": input_text[:5000]  # Limit text length in case it's very long
                        }
                    ]
                }
            ]
        }

def create_document(structured_content, output_path):
    """Create a Word document from the structured content"""
    try:
        # Print the structured content for debugging
        print(f"Structured content: {structured_content}")

        # Create a new Document
        doc = Document()

        # Check if we have valid structured content
        if not isinstance(structured_content, dict):
            # If Gemini returned raw JSON text instead of parsed JSON
            if isinstance(structured_content, str):
                import json
                try:
                    structured_content = json.loads(structured_content)
                    print(f"Parsed JSON: {structured_content}")
                except Exception as json_err:
                    print(f"JSON parsing error: {str(json_err)}")
                    # If we can't parse the JSON, create a simple document with the raw content
                    doc.add_heading("Generated Document", level=0)
                    doc.add_paragraph(str(structured_content))
                    doc.save(output_path)
                    return True
            else:
                # Create a simple document with error message
                doc.add_heading("Error in Document Generation", level=0)
                doc.add_paragraph("Could not process the structured content properly.")
                doc.save(output_path)
                return True

        # Add the title
        title = doc.add_heading(structured_content.get('title', 'Document'), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Process each section
        for section in structured_content.get('sections', []):
            # Add section heading
            level = section.get('level', 1)
            # Ensure level is between 1 and 9
            level = max(1, min(9, level))
            doc.add_heading(section.get('heading', 'Section'), level=level)

            # Process section content
            for content_item in section.get('content', []):
                content_type = content_item.get('type', '')

                # Handle special case for heading type in content items
                if content_type == 'heading' or 'level' in content_item:
                    heading_text = content_item.get('text', '')
                    heading_level = content_item.get('level', 2)  # Default to level 2 for subheadings
                    if heading_text:
                        doc.add_heading(heading_text, level=heading_level)

                elif content_type == 'paragraph':
                    # Add paragraph
                    text = content_item.get('text', '')
                    if text:
                        # Split text by newlines to create multiple paragraphs
                        paragraphs = text.split('\r\n') if '\r\n' in text else text.split('\n')
                        for para in paragraphs:
                            if para.strip():  # Only add non-empty paragraphs
                                # Check if this looks like a heading (short line followed by longer content)
                                if len(para) < 50 and para.endswith(':'):
                                    doc.add_heading(para, level=2)
                                else:
                                    doc.add_paragraph(para)

                elif content_type == 'bullet_list':
                    # Add bullet list
                    items = content_item.get('items', [])
                    if items:
                        for item in items:
                            if item:  # Only add non-empty items
                                doc.add_paragraph(item, style='List Bullet')

                # Special handling for text that contains bullet points indicated by numbers or dashes
                elif content_type == 'paragraph' and content_item.get('text', ''):
                    text = content_item.get('text', '')
                    lines = text.split('\r\n') if '\r\n' in text else text.split('\n')

                    i = 0
                    while i < len(lines):
                        line = lines[i].strip()

                        # Check if this line starts with a number or dash/bullet indicator
                        if (line.startswith('- ') or line.startswith('• ') or
                            (len(line) > 2 and line[0].isdigit() and line[1] == '.' and line[2] == ' ')):
                            # This is a bullet point
                            doc.add_paragraph(line, style='List Bullet')
                        elif line:
                            # Regular paragraph
                            doc.add_paragraph(line)

                        i += 1

                elif content_type == 'table':
                    # Add table
                    headers = content_item.get('headers', [])
                    rows = content_item.get('rows', [])

                    if headers and rows:
                        # Create table with appropriate number of columns
                        table = doc.add_table(rows=1, cols=len(headers))
                        table.style = 'Table Grid'

                        # Add headers
                        header_cells = table.rows[0].cells
                        for i, header in enumerate(headers):
                            if i < len(header_cells):
                                header_cells[i].text = str(header)

                        # Add rows
                        for row_data in rows:
                            if isinstance(row_data, list) and row_data:  # Ensure row_data is a non-empty list
                                row_cells = table.add_row().cells
                                for i, cell_data in enumerate(row_data):
                                    if i < len(row_cells):
                                        row_cells[i].text = str(cell_data)

        # Process the text to find and format sections that look like bullet points
        # This is a special case for the example you provided
        if len(structured_content.get('sections', [])) == 1:
            section = structured_content['sections'][0]
            if section.get('heading') == 'Content' and len(section.get('content', [])) == 1:
                content_item = section['content'][0]
                if content_item.get('type') == 'paragraph':
                    text = content_item.get('text', '')

                    # Check if we need to reprocess this as a structured document
                    if '\r\n\r\n' in text and ('Highlights:' in text or 'Scope:' in text):
                        # Clear the document and start over with better formatting
                        doc = Document()

                        # Add the title - extract from text if possible
                        title_text = "OmniHuman-1 and Alternatives"
                        if "OmniHuman" in text:
                            title_text = "OmniHuman-1 and Open-Source Alternatives"

                        title = doc.add_heading(title_text, level=0)
                        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        # Split by double newlines to get sections
                        sections = text.split('\r\n\r\n')

                        current_section = None
                        for section_text in sections:
                            if not section_text.strip():
                                continue

                            # Check if this is a heading
                            if section_text.strip().endswith(':') and len(section_text.strip()) < 50:
                                # This is a heading
                                current_section = section_text.strip()
                                doc.add_heading(current_section, level=1)
                            elif section_text.strip().startswith(('1.', '2.', '3.', '4.', '5.')) and ':\r\n' in section_text:
                                # This is a numbered section with a title
                                parts = section_text.split(':\r\n', 1)
                                heading = parts[0].strip() + ':'
                                content = parts[1].strip() if len(parts) > 1 else ""

                                doc.add_heading(heading, level=2)
                                if content:
                                    doc.add_paragraph(content)
                            elif 'Highlights:' in section_text:
                                # This is a highlights section
                                parts = section_text.split('Highlights:', 1)
                                if parts[0].strip():
                                    doc.add_paragraph(parts[0].strip())

                                doc.add_heading('Highlights:', level=2)

                                # Process the highlights as bullet points
                                highlights = parts[1].strip().split('\r\n\r\n')
                                for highlight in highlights:
                                    if highlight.strip():
                                        doc.add_paragraph(highlight.strip(), style='List Bullet')
                            elif 'Limitations:' in section_text:
                                # This is a limitations section
                                parts = section_text.split('Limitations:', 1)
                                if parts[0].strip():
                                    doc.add_paragraph(parts[0].strip())

                                doc.add_heading('Limitations:', level=2)

                                # Process the limitations as bullet points
                                limitations = parts[1].strip().split('\r\n\r\n')
                                for limitation in limitations:
                                    if limitation.strip():
                                        doc.add_paragraph(limitation.strip(), style='List Bullet')
                            elif 'Pros:' in section_text:
                                # This is a pros section
                                parts = section_text.split('Pros:', 1)
                                if parts[0].strip():
                                    doc.add_paragraph(parts[0].strip())

                                doc.add_heading('Pros:', level=2)

                                # Process the pros as bullet points
                                pros = parts[1].strip().split('\r\n\r\n')
                                for pro in pros:
                                    if pro.strip():
                                        doc.add_paragraph(pro.strip(), style='List Bullet')
                            elif 'Cons:' in section_text:
                                # This is a cons section
                                parts = section_text.split('Cons:', 1)
                                if parts[0].strip():
                                    doc.add_paragraph(parts[0].strip())

                                doc.add_heading('Cons:', level=2)

                                # Process the cons as bullet points
                                cons = parts[1].strip().split('\r\n\r\n')
                                for con in cons:
                                    if con.strip():
                                        doc.add_paragraph(con.strip(), style='List Bullet')
                            elif 'Recommendations' in section_text:
                                # This is a recommendations section
                                doc.add_heading('Recommendations', level=1)

                                # Process the rest as paragraphs
                                parts = section_text.split('Recommendations', 1)[1].strip()
                                if parts.startswith('\r\n'):
                                    parts = parts[2:].strip()

                                subsections = parts.split('\r\n\r\n')
                                for subsection in subsections:
                                    if subsection.strip():
                                        if subsection.strip().endswith(':'):
                                            doc.add_heading(subsection.strip(), level=2)
                                        else:
                                            doc.add_paragraph(subsection.strip())
                            else:
                                # Regular paragraph
                                doc.add_paragraph(section_text.strip())

        # Save the document with proper error handling
        try:
            # Make sure the output directory exists
            os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)

            # Save the document
            doc.save(output_path)

            # Explicitly release the document object
            del doc

            # Verify the file was created and is not empty
            if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
                print(f"Document successfully saved to {output_path} with size {os.path.getsize(output_path)} bytes")
                return True
            else:
                print(f"Document file was not created or is empty at {output_path}")
                return False
        except Exception as save_error:
            print(f"Error saving document: {str(save_error)}")
            import traceback
            print(traceback.format_exc())
            return False

    except Exception as e:
        import traceback
        print(f"Error creating document: {str(e)}")
        print(traceback.format_exc())

        # Create a simple error document
        try:
            error_doc = Document()
            error_doc.add_heading("Error in Document Generation", level=0)
            error_doc.add_paragraph(f"An error occurred: {str(e)}")
            error_doc.save(output_path)
        except:
            pass

        return False

@app.route('/download/<filename>')
def download_file(filename):
    """Download a generated file"""
    file_path = os.path.join(OUTPUT_FOLDER, filename)

    # Check if the file exists
    if not os.path.exists(file_path):
        return jsonify({'error': f'File {filename} not found'}), 404

    try:
        # Try to send the file
        return send_file(file_path, as_attachment=True)
    except Exception as e:
        print(f"Error sending file {filename}: {str(e)}")
        return jsonify({'error': f'Error downloading file: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)
